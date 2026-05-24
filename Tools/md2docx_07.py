# -*- coding: utf-8 -*-
"""
把 Plans/Dolphin设计/07_HUD组件化设计_交互与落地.md 转成同名 .docx。

目标阅读体验：
  - H1/H2/H3/H4 用 Word 内置 Heading 1~4，自动出现在导航窗格
  - 段落正文走「正文 (Body Text)」样式，11pt 中文优先字体
  - 引用块（>） 灰色斜体 + 左竖线视觉
  - 列表（- / 1.）按层级缩进，缩进每级 0.4 cm
  - 代码块（```）等宽字体 + 浅灰底纹
  - 行内 `code` 等宽字体
  - 加粗 **xx** 走 Bold；行内斜体不强求
  - 表格统一边框、首行加粗，整体居中
"""

from __future__ import annotations
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────
# 路径 / 输入输出
# ─────────────────────────────────────────────────────────────

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "Plans" / "Dolphin设计" / "07_HUD组件化设计_交互与落地.md"
DST = ROOT / "Plans" / "Dolphin设计" / "07_HUD组件化设计_交互与落地.docx"


# ─────────────────────────────────────────────────────────────
# 字体 / 颜色 / 缩进常量
# ─────────────────────────────────────────────────────────────

FONT_HAN = "Microsoft YaHei"     # 中文
FONT_LATIN = "Calibri"           # 英文 / 数字
FONT_MONO = "Consolas"           # 等宽
SIZE_BODY = Pt(11)
SIZE_QUOTE = Pt(10.5)
SIZE_CODE = Pt(10)
COLOR_QUOTE = RGBColor(0x55, 0x55, 0x55)
COLOR_INLINE_CODE = RGBColor(0xC7, 0x25, 0x4E)
COLOR_HEADING_1 = RGBColor(0x1F, 0x3A, 0x68)
COLOR_HEADING_2 = RGBColor(0x2C, 0x55, 0x8B)
COLOR_HEADING_3 = RGBColor(0x3D, 0x6E, 0xAB)
COLOR_HEADING_4 = RGBColor(0x55, 0x7A, 0xA8)


# ─────────────────────────────────────────────────────────────
# Helper：低层 XML
# ─────────────────────────────────────────────────────────────

def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_paragraph_shading(paragraph, fill_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    p_pr.append(shd)


def _set_paragraph_left_border(paragraph, color_hex: str = "888888", size_pt: int = 16) -> None:
    """给段落加左边竖线（用作 quote 视觉）。"""
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_pt))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex)
    pbdr.append(left)
    p_pr.append(pbdr)


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "B0B7C0")
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)


def _apply_run_font(run, size=SIZE_BODY, mono=False, bold=False, italic=False, color=None) -> None:
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if mono:
        run.font.name = FONT_MONO
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), FONT_MONO)
        rFonts.set(qn("w:hAnsi"), FONT_MONO)
        rFonts.set(qn("w:eastAsia"), FONT_MONO)
    else:
        run.font.name = FONT_LATIN
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), FONT_LATIN)
        rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        rFonts.set(qn("w:eastAsia"), FONT_HAN)


# ─────────────────────────────────────────────────────────────
# 行内格式：解析 **bold** / `code`
# ─────────────────────────────────────────────────────────────

# 优先匹配 `code`，再 **bold**
INLINE_PATTERN = re.compile(r"(`[^`\n]+`)|(\*\*[^*\n]+\*\*)")


def add_inline_runs(paragraph, text: str, *, base_size=SIZE_BODY, base_color=None) -> None:
    """把含 `code` / **bold** 的字符串拆段，按格式塞 run。"""
    pos = 0
    for m in INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            _apply_run_font(run, size=base_size, color=base_color)
        token = m.group(0)
        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            _apply_run_font(run, size=base_size, mono=True, color=COLOR_INLINE_CODE)
        else:  # **bold**
            run = paragraph.add_run(token[2:-2])
            _apply_run_font(run, size=base_size, bold=True, color=base_color)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _apply_run_font(run, size=base_size, color=base_color)


# ─────────────────────────────────────────────────────────────
# Markdown 解析（按行 + 块状）
# ─────────────────────────────────────────────────────────────

class Block:
    pass


class Heading(Block):
    def __init__(self, level: int, text: str):
        self.level = level
        self.text = text


class Paragraph(Block):
    def __init__(self, text: str, indent_level: int = 0):
        self.text = text
        self.indent_level = indent_level


class Quote(Block):
    def __init__(self, lines: list[str]):
        self.lines = lines


class CodeBlock(Block):
    def __init__(self, lang: str, lines: list[str]):
        self.lang = lang
        self.lines = lines


class ListItem(Block):
    def __init__(self, ordered: bool, text: str, indent_level: int = 0, marker: str = ""):
        self.ordered = ordered
        self.text = text
        self.indent_level = indent_level
        self.marker = marker  # 给 ordered 保留原始数字字符串


class Table(Block):
    def __init__(self, header: list[str], rows: list[list[str]]):
        self.header = header
        self.rows = rows


class HRule(Block):
    pass


def parse_markdown(md_text: str) -> list[Block]:
    lines = md_text.splitlines()
    blocks: list[Block] = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            i += 1
            continue

        # 水平线
        if re.match(r"^-{3,}\s*$", stripped) or re.match(r"^\*{3,}\s*$", stripped):
            blocks.append(HRule())
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            blocks.append(Heading(level, text))
            i += 1
            continue

        # 代码块
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code_lines: list[str] = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append(CodeBlock(lang, code_lines))
            continue

        # 引用块（连续 > ...）
        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while i < n and lines[i].strip().startswith(">"):
                content = re.sub(r"^\s*>\s?", "", lines[i])
                quote_lines.append(content)
                i += 1
            blocks.append(Quote(quote_lines))
            continue

        # 表格（必须 |...|...| 且下一行是 ---|---）
        if "|" in stripped and i + 1 < n and re.match(r"^\s*\|?\s*[-:]+", lines[i + 1].strip()):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            rows: list[list[str]] = []
            i += 2  # 跳过 header + 分隔行
            while i < n and "|" in lines[i] and lines[i].strip():
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            blocks.append(Table(header, rows))
            continue

        # 无序列表 (- / *)
        m_ul = re.match(r"^(\s*)([-*])\s+(.*)$", line)
        if m_ul:
            indent = len(m_ul.group(1)) // 2  # 2 空格 = 一级
            text = m_ul.group(3).strip()
            blocks.append(ListItem(False, text, indent))
            i += 1
            continue

        # 有序列表 (1. / 2.)
        m_ol = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m_ol:
            indent = len(m_ol.group(1)) // 2
            text = m_ol.group(3).strip()
            blocks.append(ListItem(True, text, indent, marker=m_ol.group(2)))
            i += 1
            continue

        # 普通段落（合并连续非空、非特殊行直到空行）
        para_lines = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if not nxt:
                break
            if nxt.startswith("#") or nxt.startswith(">") or nxt.startswith("```"):
                break
            if re.match(r"^[-*]\s+", nxt) or re.match(r"^\d+\.\s+", nxt):
                break
            if "|" in nxt and i + 1 < n and re.match(r"^\s*\|?\s*[-:]+", lines[i + 1].strip()):
                break
            if re.match(r"^-{3,}\s*$", nxt) or re.match(r"^\*{3,}\s*$", nxt):
                break
            para_lines.append(nxt)
            i += 1
        blocks.append(Paragraph(" ".join(para_lines)))

    return blocks


# ─────────────────────────────────────────────────────────────
# 渲染 Block 到 docx
# ─────────────────────────────────────────────────────────────

def render(doc: Document, blocks: list[Block]) -> None:
    for blk in blocks:
        if isinstance(blk, Heading):
            render_heading(doc, blk)
        elif isinstance(blk, Paragraph):
            render_paragraph(doc, blk)
        elif isinstance(blk, Quote):
            render_quote(doc, blk)
        elif isinstance(blk, CodeBlock):
            render_code(doc, blk)
        elif isinstance(blk, ListItem):
            render_list_item(doc, blk)
        elif isinstance(blk, Table):
            render_table(doc, blk)
        elif isinstance(blk, HRule):
            render_hrule(doc)


def render_heading(doc: Document, blk: Heading) -> None:
    # 顶级 Heading 1 当作"章"标题，整篇文档第一行的 H1 也走 Heading 1
    level = max(1, min(blk.level, 4))
    p = doc.add_heading(level=level)
    run = p.add_run(blk.text)
    color_map = {1: COLOR_HEADING_1, 2: COLOR_HEADING_2, 3: COLOR_HEADING_3, 4: COLOR_HEADING_4}
    size_map = {1: Pt(20), 2: Pt(16), 3: Pt(13.5), 4: Pt(12)}
    _apply_run_font(run, size=size_map[level], bold=True, color=color_map[level])
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)


def render_paragraph(doc: Document, blk: Paragraph) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    if blk.indent_level > 0:
        p.paragraph_format.left_indent = Cm(0.4 * blk.indent_level)
    add_inline_runs(p, blk.text)


def render_quote(doc: Document, blk: Quote) -> None:
    text = "\n".join(blk.lines).strip()
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    _set_paragraph_shading(p, "F4F6FB")
    _set_paragraph_left_border(p, color_hex="6B8AC8", size_pt=18)
    add_inline_runs(p, text, base_size=SIZE_QUOTE, base_color=COLOR_QUOTE)


def render_code(doc: Document, blk: CodeBlock) -> None:
    if not blk.lines:
        return
    # 用 1 列单元格表格做代码框，方便底纹与等宽字体一起生效
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "F5F5F0")
    cell.paragraphs[0].text = ""  # 清空默认段
    for idx, line in enumerate(blk.lines):
        if idx == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        _apply_run_font(run, size=SIZE_CODE, mono=True)
    # 表格上下间距
    spacer_before = doc.add_paragraph()
    spacer_before.paragraph_format.space_after = Pt(0)
    # 移到代码块前？ python-docx 不支持往回移 element 容易，省略
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def render_list_item(doc: Document, blk: ListItem) -> None:
    # 不依赖 Word 的内置编号样式（多级编号容易跑偏）；
    # 自己拼"  • " / "  1. "前缀 + 缩进，更可控、更稳。
    indent_cm = 0.5 + 0.5 * blk.indent_level
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.first_line_indent = Cm(-0.5)  # 悬挂缩进，让标记视觉对齐
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.35

    # 前缀符号
    if blk.ordered:
        prefix = f"{blk.marker}. "
    else:
        symbols = ["•", "◦", "▪"]
        prefix = symbols[min(blk.indent_level, len(symbols) - 1)] + " "
    run = p.add_run(prefix)
    _apply_run_font(run, size=SIZE_BODY, bold=False)

    add_inline_runs(p, blk.text)


def render_table(doc: Document, blk: Table) -> None:
    cols = max(len(blk.header), max((len(r) for r in blk.rows), default=0))
    if cols == 0:
        return
    table = doc.add_table(rows=1 + len(blk.rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_borders(table)

    # Header
    hdr_row = table.rows[0]
    for c in range(cols):
        cell = hdr_row.cells[c]
        _set_cell_shading(cell, "DCE7F6")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text = blk.header[c] if c < len(blk.header) else ""
        run = p.add_run(text)
        _apply_run_font(run, size=Pt(10.5), bold=True, color=COLOR_HEADING_2)

    # Body
    for r, row in enumerate(blk.rows, start=1):
        for c in range(cols):
            cell = table.rows[r].cells[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.3
            text = row[c] if c < len(row) else ""
            add_inline_runs(p, text, base_size=Pt(10.5))

    # 表格后空一行
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def render_hrule(doc: Document) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pbdr.append(bottom)
    p_pr.append(pbdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


# ─────────────────────────────────────────────────────────────
# 全局样式（页边距 / 默认正文字体）
# ─────────────────────────────────────────────────────────────

def setup_document(doc: Document) -> None:
    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.0)

    # Normal 样式默认字体
    style = doc.styles["Normal"]
    style.font.name = FONT_LATIN
    style.font.size = SIZE_BODY
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_LATIN)
    rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    rFonts.set(qn("w:eastAsia"), FONT_HAN)


# ─────────────────────────────────────────────────────────────
# main
# ─────────────────────────────────────────────────────────────

def main() -> int:
    if not SRC.exists():
        print(f"[ERROR] 源文件不存在：{SRC}", file=sys.stderr)
        return 1

    md_text = SRC.read_text(encoding="utf-8")
    blocks = parse_markdown(md_text)

    doc = Document()
    setup_document(doc)
    render(doc, blocks)

    DST.parent.mkdir(parents=True, exist_ok=True)
    target = DST
    try:
        doc.save(str(target))
    except PermissionError:
        # 目标被 Word 等程序占用 → 写到带时间戳的兜底副本
        from datetime import datetime
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        target = DST.with_name(f"{DST.stem}_{stamp}{DST.suffix}")
        doc.save(str(target))
        print(f"[WARN] 主文件被占用，已改写到副本：{target.name}")
    print(f"[OK] 已生成：{target}  ({target.stat().st_size / 1024:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
